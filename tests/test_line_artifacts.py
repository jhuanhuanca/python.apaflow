"""Tests de eliminación de líneas decorativas ChatGPT."""

import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.apa7.ai_sanitizer import SanitizeStats
from app.apa7.engine import format_docx_apa7
from app.apa7.line_artifacts import _SEPARATOR_TEXT_RE, remove_chatgpt_horizontal_lines


def _add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def test_separator_text_pattern():
    assert _SEPARATOR_TEXT_RE.match("──────────────")
    assert _SEPARATOR_TEXT_RE.match("___")
    assert not _SEPARATOR_TEXT_RE.match("Texto normal")


def test_removes_empty_paragraph_with_border():
    doc = Document()
    doc.add_paragraph("2. ARQUITECTURA DEL SISTEMA")
    empty = doc.add_paragraph("")
    _add_bottom_border(empty)
    doc.add_paragraph("Objetivo")

    stats = SanitizeStats()
    remove_chatgpt_horizontal_lines(doc, stats)

    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "2. ARQUITECTURA DEL SISTEMA" in texts
    assert "Objetivo" in texts
    assert stats.decorative_paragraphs_removed >= 1


def test_removes_separator_only_paragraph():
    doc = Document()
    doc.add_paragraph("Intro")
    doc.add_paragraph("────────────────")
    doc.add_paragraph("Contenido")

    stats = SanitizeStats()
    remove_chatgpt_horizontal_lines(doc, stats)

    assert "────────────────" not in [p.text for p in doc.paragraphs]
    assert stats.decorative_paragraphs_removed >= 1


def test_full_pipeline_removes_lines():
    doc = Document()
    doc.add_paragraph("Sección")
    doc.add_paragraph("___")
    buf = io.BytesIO()
    doc.save(buf)

    processed, report = format_docx_apa7(buf.getvalue())
    out = Document(io.BytesIO(processed))
    assert "___" not in [p.text.strip() for p in out.paragraphs]
    assert report["ai_sanitize"]["decorative_paragraphs_removed"] >= 1
