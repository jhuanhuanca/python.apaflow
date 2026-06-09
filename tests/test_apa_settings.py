"""Tests de configuración APA personalizada."""

import io
import json

from docx import Document
from docx.shared import Inches

from app.apa7.engine import format_docx_apa7
from app.apa7.settings import ApaRuntimeSettings


def test_custom_margins_applied():
    doc = Document()
    doc.add_paragraph("Contenido de prueba")
    buf = io.BytesIO()
    doc.save(buf)

    settings = ApaRuntimeSettings.from_payload(
        {"margins_in": {"top": 1.5, "bottom": 1.5, "left": 1.2, "right": 1.2}}
    )
    processed, report = format_docx_apa7(buf.getvalue(), apa_settings=settings)
    out = Document(io.BytesIO(processed))
    section = out.sections[0]
    assert abs(section.top_margin.inches - 1.5) < 0.01
    assert abs(section.left_margin.inches - 1.2) < 0.01
    assert report["apa_settings"]["margins_in"]["top"] == 1.5


def test_parse_json_settings():
    raw = json.dumps({"line_spacing": 1.5, "font_size_pt": 11})
    settings = ApaRuntimeSettings.parse_json(raw)
    assert settings.line_spacing == 1.5
    assert settings.font_size_pt == 11
