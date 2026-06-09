"""Orquestación del formateo APA 7 y informe de cambios."""

from __future__ import annotations

import io
import json
from dataclasses import asdict, dataclass, field
from typing import Any

from docx import Document

from app.apa7.ai_sanitizer import sanitize_document
from app.apa7.classifier import ParagraphClassifier
from app.apa7.formatter import (
    apply_block_quote_style,
    apply_body_style,
    apply_figure_caption,
    apply_heading,
    apply_reference_entry,
    apply_refs_heading,
    apply_table_caption,
    format_tables,
    insert_toc_at_start,
    set_apa_page_setup,
    set_header_page_number,
)
from app.apa7.oxml_helpers import document_has_toc_field
from app.apa7.settings import ApaRuntimeSettings


@dataclass
class FormatReport:
    paragraphs_processed: int = 0
    headings: dict[str, int] = field(default_factory=dict)
    reference_entries: int = 0
    block_quotes: int = 0
    table_captions: int = 0
    figure_captions: int = 0
    tables_formatted: int = 0
    toc_inserted: bool = False
    page_numbers: str = "header_right"
    apa_settings: dict[str, Any] = field(default_factory=dict)
    ai_sanitize: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)

    def bump_heading(self, level: int) -> None:
        key = f"h{level}"
        self.headings[key] = self.headings.get(key, 0) + 1

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _paragraph_style_name(paragraph) -> str | None:
    try:
        return paragraph.style.name if paragraph.style else None
    except (AttributeError, KeyError):
        return None


def format_docx_apa7(
    content: bytes,
    apa_settings: ApaRuntimeSettings | dict[str, Any] | None = None,
) -> tuple[bytes, dict[str, Any]]:
    """
    Formatea un .docx según APA 7 (edición estudiantil).

    Args:
        content: bytes del .docx original.
        apa_settings: configuración personalizada (Pro) o dict JSON-compatible.

    Returns:
        Tupla (bytes del documento, informe JSON-serializable).
    """
    if isinstance(apa_settings, dict):
        settings = ApaRuntimeSettings.from_payload(apa_settings)
    elif isinstance(apa_settings, ApaRuntimeSettings):
        settings = apa_settings
    else:
        settings = ApaRuntimeSettings()

    doc = Document(io.BytesIO(content))
    report = FormatReport(
        page_numbers=settings.page_number_position,
        apa_settings=settings.to_dict(),
    )
    classifier = ParagraphClassifier()

    sanitize_stats = sanitize_document(doc)
    report.ai_sanitize = sanitize_stats.to_dict()
    if sanitize_stats.total_chars_removed > 0 or sanitize_stats.ai_horizontal_lines_removed > 0:
        parts = []
        if sanitize_stats.total_chars_removed > 0:
            parts.append(f"{sanitize_stats.total_chars_removed} caracteres/emojis IA")
        if sanitize_stats.ai_horizontal_lines_removed > 0 or sanitize_stats.decorative_paragraphs_removed > 0:
            lines = sanitize_stats.ai_horizontal_lines_removed + sanitize_stats.decorative_paragraphs_removed
            parts.append(f"{lines} líneas decorativas ChatGPT")
        report.warnings.append("Se eliminaron artefactos IA: " + ", ".join(parts) + ".")

    set_apa_page_setup(doc, settings)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        stripped = text.strip()
        if not stripped:
            continue

        kind = classifier.classify(stripped, _paragraph_style_name(paragraph))
        report.paragraphs_processed += 1

        if kind == "refs_heading":
            apply_refs_heading(paragraph, settings)
            report.bump_heading(1)
        elif kind == "reference_entry":
            apply_reference_entry(paragraph, settings)
            report.reference_entries += 1
        elif kind == "block_quote":
            apply_block_quote_style(paragraph, settings)
            report.block_quotes += 1
        elif kind == "table_caption":
            apply_table_caption(paragraph, stripped, settings)
            report.table_captions += 1
        elif kind == "figure_caption":
            apply_figure_caption(paragraph, stripped, settings)
            report.figure_captions += 1
        elif kind.startswith("h") and kind[1:].isdigit():
            level = int(kind[1:])
            apply_heading(paragraph, level, settings)
            report.bump_heading(level)
        else:
            apply_body_style(paragraph, settings)

    report.tables_formatted = format_tables(doc, settings)

    if not document_has_toc_field(doc):
        insert_toc_at_start(doc, settings=settings)
        report.toc_inserted = True
    else:
        report.warnings.append("Ya existía un campo TOC; no se insertó otro índice.")

    for section in doc.sections:
        set_header_page_number(section, settings)

    if not doc.paragraphs:
        report.warnings.append("Documento sin párrafos de texto.")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), report.to_dict()


def report_to_json(report: dict[str, Any]) -> str:
    return json.dumps(report, ensure_ascii=False)
