"""
Elimina emojis y caracteres invisibles típicos de textos copiados desde IA.

Incluye marcadores de huella digital (espacios de ancho cero, tags Unicode,
selectores de variación, operadores invisibles) que algunas herramientas insertan
para rastreo o formato decorativo no apto para trabajos académicos APA.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import asdict, dataclass, field
from typing import Any

from docx.document import Document
from docx.text.paragraph import Paragraph

# --- Código Unicode explícito (huellas / control / formato invisible) ---
_INVISIBLE_CODEPOINTS = frozenset(
    {
        0x00AD,  # guión blando
        0x034F,  # combining grapheme joiner
        0x061C,  # arabic letter mark
        0x115F,  # hangul chaeum
        0x1160,  # hangul filler
        0x17B4,  # khmer vowel inherent
        0x17B5,
        0x180E,  # mongolian vowel separator (deprecated)
        0x200B,  # zero width space
        0x200C,  # zero width non-joiner
        0x200D,  # zero width joiner
        0x200E,  # left-to-right mark
        0x200F,  # right-to-left mark
        0x202A,
        0x202B,
        0x202C,
        0x202D,
        0x202E,  # bidi overrides
        0x2060,
        0x2061,
        0x2062,
        0x2063,
        0x2064,  # invisible operators / separator
        0x2066,
        0x2067,
        0x2068,
        0x2069,  # bidi isolates
        0x3164,  # hangul filler
        0xFEFF,  # BOM / zero width no-break space
        0xFFFC,  # object replacement
        0xFFF9,
        0xFFFA,
        0xFFFB,  # interlinear annotation anchors
    }
)

# Tags Unicode (U+E0001–U+E007F): a veces usados como metadatos ocultos
_UNICODE_TAG_RE = re.compile(r"[\U000E0001-\U000E007F]")

# Selectores de variación + modificadores Fitzpatrick
_VARIATION_AND_MODIFIERS_RE = re.compile(
    "["
    "\uFE00-\uFE0F"
    "\U0001F3FB-\U0001F3FF"
    "]",
    flags=re.UNICODE,
)

# Emojis y pictogramas (bloques principales)
_EMOJI_RE = re.compile(
    "["
    "\U0001F1E0-\U0001F1FF"
    "\U0001F300-\U0001F5FF"
    "\U0001F600-\U0001F64F"
    "\U0001F680-\U0001F6FF"
    "\U0001F700-\U0001F77F"
    "\U0001F780-\U0001F7FF"
    "\U0001F800-\U0001F8FF"
    "\U0001F900-\U0001F9FF"
    "\U0001FA00-\U0001FAFF"
    "\U00002600-\U000026FF"
    "\U00002700-\U000027BF"
    "\U0001F000-\U0001F02F"
    "\U0001FA70-\U0001FAFF"
    "]+",
    flags=re.UNICODE,
)

# Símbolos decorativos frecuentes en respuestas de chatbots (no académicos)
_DECORATIVE_SYMBOLS_RE = re.compile(
    r"[\u2728\u2705\u274C\u2757\u2753\u2B50\u2729\u2734\u2733\u2611\u2610\u2714\u2716\u2718]"
)

# Espacios Unicode raros → espacio normal
_SPECIAL_SPACES = str.maketrans(
    {
        "\u00a0": " ",  # NBSP
        "\u2000": " ",
        "\u2001": " ",
        "\u2002": " ",
        "\u2003": " ",
        "\u2004": " ",
        "\u2005": " ",
        "\u2006": " ",
        "\u2007": " ",
        "\u2008": " ",
        "\u2009": " ",
        "\u200a": " ",
        "\u202f": " ",
        "\u205f": " ",
        "\u3000": " ",
    }
)

# Frases de descargo típicas de IA (párrafo completo o línea)
_AI_DISCLAIMER_RE = re.compile(
    r"^\s*(?:"
    r"como\s+(?:modelo\s+de\s+lenguaje\s+de\s+ia|inteligencia\s+artificial)"
    r"|as\s+an\s+ai\s+(?:language\s+)?model"
    r"|i(?:'m|\s+am)\s+an?\s+ai"
    r"|no\s+puedo\s+(?:acceder|browse)"
    r"|chatgpt|openai|claude|gemini|copilot"
    r").*$",
    re.IGNORECASE | re.MULTILINE,
)

# Línea decorativa solo con guiones/símbolos (markdown ChatGPT pegado como texto)
_SEPARATOR_LINE_TEXT_RE = re.compile(
    r"^[\s\-–—_=~·•─━│┃]{2,}$",
    re.UNICODE,
)


@dataclass
class SanitizeStats:
    emojis_removed: int = 0
    invisible_chars_removed: int = 0
    unicode_tags_removed: int = 0
    decorative_symbols_removed: int = 0
    ai_disclaimer_lines_removed: int = 0
    runs_touched: int = 0
    paragraphs_touched: int = 0
    headers_footers_touched: int = 0
    table_cells_touched: int = 0
    core_properties_cleaned: int = 0
    ai_horizontal_lines_removed: int = 0
    decorative_paragraphs_removed: int = 0
    paragraph_borders_removed: int = 0

    def merge(self, other: SanitizeStats) -> None:
        for key in self.__dataclass_fields__:
            setattr(self, key, getattr(self, key) + getattr(other, key))

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @property
    def total_chars_removed(self) -> int:
        return (
            self.emojis_removed
            + self.invisible_chars_removed
            + self.unicode_tags_removed
            + self.decorative_symbols_removed
        )


def _count_removed(before: str, after: str) -> int:
    return max(0, len(before) - len(after))


def _strip_invisible_codepoints(text: str) -> tuple[str, int]:
    removed = 0
    chars: list[str] = []
    for ch in text:
        if ord(ch) in _INVISIBLE_CODEPOINTS:
            removed += 1
            continue
        chars.append(ch)
    return "".join(chars), removed


def sanitize_plain_text(text: str) -> tuple[str, SanitizeStats]:
    """Limpia una cadena sin modificar el documento Word."""
    stats = SanitizeStats()
    if not text:
        return text, stats

    result = text.translate(_SPECIAL_SPACES)

    result, n = _strip_invisible_codepoints(result)
    stats.invisible_chars_removed += n

    new_result, n_tags = _UNICODE_TAG_RE.subn("", result)
    stats.unicode_tags_removed += n_tags
    result = new_result

    new_result, n_var = _VARIATION_AND_MODIFIERS_RE.subn("", result)
    stats.invisible_chars_removed += n_var
    result = new_result

    new_result, n_emoji = _EMOJI_RE.subn("", result)
    stats.emojis_removed += n_emoji
    result = new_result

    new_result, n_deco = _DECORATIVE_SYMBOLS_RE.subn("", result)
    stats.decorative_symbols_removed += n_deco
    result = new_result

    # Categoría So pictográfica residual (emoji sin match en bloques)
    filtered: list[str] = []
    for ch in result:
        if _is_pictographic_symbol(ch):
            stats.emojis_removed += 1
            continue
        filtered.append(ch)
    result = "".join(filtered)

    result = re.sub(r"[ \t]+", " ", result)
    result = re.sub(r" ?\n ?", "\n", result)

    if _SEPARATOR_LINE_TEXT_RE.match(result.strip()):
        stats.decorative_symbols_removed += len(result.strip())
        return "", stats

    return result, stats


def _is_pictographic_symbol(ch: str) -> bool:
    if len(ch) != 1:
        return False
    cp = ord(ch)
    if cp in _INVISIBLE_CODEPOINTS:
        return True
    name = unicodedata.name(ch, "")
    if "EMOJI" in name or "EMOTICON" in name:
        return True
    cat = unicodedata.category(ch)
    if cat == "So" and cp >= 0x1F000:
        return True
    return False


def _sanitize_runs_in_paragraph(paragraph: Paragraph) -> SanitizeStats:
    stats = SanitizeStats()
    for run in paragraph.runs:
        original = run.text
        if not original:
            continue
        cleaned, run_stats = sanitize_plain_text(original)
        if cleaned != original:
            run.text = cleaned
            stats.merge(run_stats)
            stats.runs_touched += 1
    if stats.runs_touched:
        stats.paragraphs_touched += 1
    return stats


def _remove_ai_disclaimer_paragraphs(doc: Document, stats: SanitizeStats) -> None:
    """Elimina párrafos que son solo avisos de IA (no contenido académico)."""
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if _AI_DISCLAIMER_RE.match(text):
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)
                stats.ai_disclaimer_lines_removed += 1


def _sanitize_core_properties(doc: Document, stats: SanitizeStats) -> None:
    """Quita menciones obvias a herramientas IA en metadatos editables."""
    props = doc.core_properties
    ai_markers = ("chatgpt", "openai", "claude", "gemini", "copilot", "gpt-4", "gpt-3")

    for attr in ("comments", "keywords", "subject", "category"):
        value = getattr(props, attr, None)
        if not value or not isinstance(value, str):
            continue
        lowered = value.lower()
        if any(marker in lowered for marker in ai_markers):
            setattr(props, attr, "")
            stats.core_properties_cleaned += 1


def sanitize_document(doc: Document) -> SanitizeStats:
    """
    Recorre cuerpo, tablas, encabezados y pies; elimina artefactos de IA.

    Se ejecuta antes del formateo APA para no conservar huellas invisibles.
    """
    total = SanitizeStats()

    from app.apa7.line_artifacts import remove_chatgpt_horizontal_lines

    # 1) Líneas/bordes ChatGPT (incluye párrafos "---" antes de limpiar texto)
    remove_chatgpt_horizontal_lines(doc, total)

    for paragraph in doc.paragraphs:
        total.merge(_sanitize_runs_in_paragraph(paragraph))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cell_stats = _sanitize_runs_in_paragraph(paragraph)
                    if cell_stats.runs_touched:
                        total.table_cells_touched += 1
                    total.merge(cell_stats)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                hf_stats = _sanitize_runs_in_paragraph(paragraph)
                if hf_stats.runs_touched:
                    total.headers_footers_touched += 1
                total.merge(hf_stats)

    _remove_ai_disclaimer_paragraphs(doc, total)
    _sanitize_core_properties(doc, total)

    # 2) Segunda pasada: bordes vacíos tras quitar emojis/símbolos en runs
    remove_chatgpt_horizontal_lines(doc, total)

    return total
