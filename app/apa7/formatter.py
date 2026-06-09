"""Aplicación de estilos APA 7 a párrafos, tablas y secciones."""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from app.apa7.constants import (
    BLOCK_QUOTE_INDENT_INCH,
    FIGURE_CAPTION_PATTERN,
    HANGING_INDENT_INCH,
    HEADING_INDENT_INCH,
    TABLE_CAPTION_PATTERN,
    TABLE_FONT_SIZE_PT,
)
from app.apa7.oxml_helpers import add_field_run, clear_paragraph_runs, set_run_font
from app.apa7.settings import ApaRuntimeSettings


def _font_kwargs(settings: ApaRuntimeSettings) -> dict:
    return {"font_name": settings.font_name, "size_pt": settings.font_size_pt}


def _apply_line_spacing(fmt, settings: ApaRuntimeSettings) -> None:
    spacing = settings.line_spacing
    if spacing <= 1.05:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif spacing <= 1.6:
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if spacing != 2.0:
            fmt.line_spacing = spacing


def set_apa_page_setup(doc: Document, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(cfg.margin_top)
        section.bottom_margin = Inches(cfg.margin_bottom)
        section.left_margin = Inches(cfg.margin_left)
        section.right_margin = Inches(cfg.margin_right)


def set_header_page_number(section, settings: ApaRuntimeSettings | None = None) -> None:
    """Número de página según configuración (header derecha o footer centrado)."""
    cfg = settings or ApaRuntimeSettings()
    header = section.header
    footer = section.footer

    for paragraph in header.paragraphs:
        clear_paragraph_runs(paragraph)
    for paragraph in footer.paragraphs:
        clear_paragraph_runs(paragraph)

    if cfg.page_number_position == "footer_center":
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field_run(paragraph, "PAGE")
        return

    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field_run(paragraph, "PAGE")


def _apply_paragraph_style_base(paragraph) -> None:
    try:
        paragraph.style = paragraph.part.document.styles["Normal"]
    except KeyError:
        pass


def apply_body_style(paragraph, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    _apply_paragraph_style_base(paragraph)
    font = _font_kwargs(cfg)
    for run in paragraph.runs:
        set_run_font(run, **font)
    fmt = paragraph.paragraph_format
    _apply_line_spacing(fmt, cfg)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Inches(0)
    fmt.first_line_indent = Inches(0.5)
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_block_quote_style(paragraph, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    _apply_paragraph_style_base(paragraph)
    font = _font_kwargs(cfg)
    for run in paragraph.runs:
        set_run_font(run, **font)
    fmt = paragraph.paragraph_format
    _apply_line_spacing(fmt, cfg)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Inches(BLOCK_QUOTE_INDENT_INCH)
    fmt.right_indent = Inches(BLOCK_QUOTE_INDENT_INCH)
    fmt.first_line_indent = Inches(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_heading(paragraph, level: int, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    style_name = f"Heading {level}"
    try:
        paragraph.style = paragraph.part.document.styles[style_name]
    except KeyError:
        _apply_paragraph_style_base(paragraph)

    font = _font_kwargs(cfg)
    for run in paragraph.runs:
        italic = level in (3, 5)
        set_run_font(run, bold=True, italic=italic, **font)

    fmt = paragraph.paragraph_format
    _apply_line_spacing(fmt, cfg)
    fmt.space_before = Pt(12 if level == 1 else 6)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Inches(0)
    fmt.right_indent = Inches(0)

    if level == 1:
        fmt.left_indent = Inches(0)
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level in (4, 5):
        fmt.left_indent = Inches(HEADING_INDENT_INCH)
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        fmt.left_indent = Inches(0)
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_refs_heading(paragraph, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    apply_heading(paragraph, 1, cfg)
    font = _font_kwargs(cfg)
    for run in paragraph.runs:
        set_run_font(run, bold=True, **font)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def apply_reference_entry(paragraph, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    _apply_paragraph_style_base(paragraph)
    font = _font_kwargs(cfg)
    for run in paragraph.runs:
        set_run_font(run, **font)
    fmt = paragraph.paragraph_format
    _apply_line_spacing(fmt, cfg)
    fmt.left_indent = Inches(HANGING_INDENT_INCH)
    fmt.first_line_indent = Inches(-HANGING_INDENT_INCH)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_table_caption(paragraph, text: str, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    match = TABLE_CAPTION_PATTERN.match(text.strip())
    if not match:
        apply_body_style(paragraph, cfg)
        return

    label = f"{match.group(1).capitalize()} {match.group(2)}"
    title = (match.group(3) or "").strip()

    clear_paragraph_runs(paragraph)
    _apply_paragraph_style_base(paragraph)
    fmt = paragraph.paragraph_format
    _apply_line_spacing(fmt, cfg)
    fmt.first_line_indent = Inches(0)
    fmt.left_indent = Inches(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    font = _font_kwargs(cfg)
    run_label = paragraph.add_run(f"{label}\n")
    set_run_font(run_label, bold=True, **font)
    if title:
        run_title = paragraph.add_run(title)
        set_run_font(run_title, italic=True, **font)


def apply_figure_caption(paragraph, text: str, settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    match = FIGURE_CAPTION_PATTERN.match(text.strip())
    if not match:
        apply_body_style(paragraph, cfg)
        return

    label = f"{match.group(1).capitalize()} {match.group(2)}."
    title = (match.group(3) or "").strip()

    clear_paragraph_runs(paragraph)
    _apply_paragraph_style_base(paragraph)
    fmt = paragraph.paragraph_format
    _apply_line_spacing(fmt, cfg)
    fmt.first_line_indent = Inches(0)
    fmt.left_indent = Inches(0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    font = _font_kwargs(cfg)
    run_label = paragraph.add_run(f"{label} ")
    set_run_font(run_label, bold=True, **font)
    if title:
        run_title = paragraph.add_run(title)
        set_run_font(run_title, italic=True, **font)


def format_tables(doc: Document, settings: ApaRuntimeSettings | None = None) -> int:
    """Tipografía en celdas; interlineado sencillo (permitido en tablas APA)."""
    cfg = settings or ApaRuntimeSettings()
    font = _font_kwargs(cfg)
    count = 0
    for table in doc.tables:
        count += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    _apply_paragraph_style_base(paragraph)
                    for run in paragraph.runs:
                        set_run_font(run, size_pt=TABLE_FONT_SIZE_PT, font_name=cfg.font_name)
                    fmt = paragraph.paragraph_format
                    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.first_line_indent = Inches(0)
                    fmt.left_indent = Inches(0)
                    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return count


def insert_toc_at_start(doc: Document, *, title: str = "Índice", settings: ApaRuntimeSettings | None = None) -> None:
    cfg = settings or ApaRuntimeSettings()
    body = doc.element.body

    title_p = doc.add_paragraph(title)
    body.remove(title_p._p)
    body.insert(0, title_p._p)
    apply_heading(title_p, 1, cfg)
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_p = doc.add_paragraph()
    body.remove(toc_p._p)
    body.insert(1, toc_p._p)
    add_field_run(toc_p, r'TOC \o "1-5" \h \z \u')
